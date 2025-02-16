import os
import openai
import torch
import whisper
from dotenv import load_dotenv
from docx import Document

# ✅ 환경 변수 로드
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

# ✅ macOS 다운로드 폴더 설정
download_path = os.path.join(os.path.expanduser('~'), 'Downloads')

# ✅ GPU 지원 여부 확인
device = "cuda" if torch.cuda.is_available() else "cpu"
print(f"🖥 Using device: {device}")

# ✅ Whisper 모델 로드
model = whisper.load_model("medium").to(device)
print("Whisper model loaded successfully!")

def download_youtube_audio(youtube_link, output_format='mp3'):
    """ YouTube 오디오 다운로드 및 변환 """
    command = f'yt-dlp -x --audio-format {output_format} -o "{download_path}/%(title)s.{output_format}" {youtube_link}'
    print(f"Executing command: {command}")
    
    result = os.system(command)
    if result == 0:
        # 다운로드된 파일명 확인
        files = [f for f in os.listdir(download_path) if f.endswith(f".{output_format}")]
        if files:
            audio_file = os.path.join(download_path, files[-1])  # 최신 파일 반환
            print(f"✅ 오디오 다운로드 성공: {audio_file}")
            return audio_file
        else:
            print("❌ 다운로드된 MP3 파일을 찾을 수 없습니다.")
            return None
    else:
        print("❌ 오디오 다운로드 및 변환 실패")
        return None

def transcribe_and_translate(audio_path):
    """ Whisper를 이용한 오디오 전사 및 OpenAI API 번역 """
    print(f"🎤 Transcribing audio file: {audio_path}")

    if not os.path.exists(audio_path):
        print("❌ Audio file not found, skipping transcription.")
        return None, None

    try:
        print("⚡ Whisper 모델이 전사 중... (이 과정이 오래 걸릴 수 있음)")
        transcript = model.transcribe(audio_path, language="ja", initial_prompt="日本語の音声をテキストに変換してください。")
        print("✅ 전사 완료!") 

        print(f"📝 Whisper Output: {transcript}")  

        japanese_transcript = [
            {"start": segment["start"], "end": segment["end"], "text": segment["text"]}
            for segment in transcript["segments"] if segment["text"]
        ]
        print(f"📜 Extracted Japanese Transcript: {japanese_transcript}")

        summary, translated_entries = summarize_and_translate(japanese_transcript)
        return summary, translated_entries

    except Exception as e:
        print("❌ Whisper 실행 중 오류 발생:", e)
        return None, None

def summarize_and_translate(transcript_data):
    """ OpenAI API를 사용하여 일본어 텍스트 요약 및 번역 """
    openai.api_key = os.getenv("OPENAI_API_KEY")
    if not transcript_data:
        print("❌ No transcript data found for translation.")
        return None, None

    full_text = " ".join([entry["text"] for entry in transcript_data])

    try:
        print("📢 Sending text to OpenAI API for summarization...")
        summary_response = openai.Completion.create(
            engine="davinci",
            prompt=f"Summarize this Japanese text: '{full_text}'",
            max_tokens=150
        )
        summary = summary_response.choices[0].text.strip()
        print(f"📝 Summary: {summary}")

        translated_entries = []
        for entry in transcript_data:
            print(f"📢 Sending text to OpenAI API for translation: {entry['text']}")
            translation_response = openai.Completion.create(
                engine="davinci",
                prompt=f"Translate this Japanese text into Korean: '{entry['text']}'",
                max_tokens=150
            )
            translated_text = translation_response.choices[0].text.strip()
            print(f"✅ Translated: {translated_text}")

            translated_entries.append({
                "start": entry["start"],
                "end": entry["end"],
                "japanese": entry["text"],
                "korean": translated_text
            })

        return summary, translated_entries
    except Exception as e:
        print("❌ OpenAI API 요청 중 오류 발생:", e)
        return None, None

def save_transcript_to_docx(summary, translated_data, output_filename="translated_transcript.docx"):
    """ 번역된 내용을 DOCX 파일로 저장 """
    output_path = os.path.join(download_path, output_filename)
    print(f"📂 Saving translated document to: {output_path}")

    doc = Document()
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(summary)
    doc.add_page_break()

    doc.add_heading('Translations', level=1)
    for entry in translated_data:
        doc.add_paragraph(f"[{entry['start']}s - {entry['end']}s]", style="Heading 3")
        doc.add_paragraph(entry["japanese"], style="Normal")
        doc.add_paragraph(entry["korean"], style="Normal")
        doc.add_paragraph("")

    doc.save(output_path)
    print(f"✅ Translated document saved successfully at: {output_path}")

if __name__ == "__main__":
    youtube_link = input("Please enter the YouTube link: ").strip()
    if youtube_link:
        audio_path = download_youtube_audio(youtube_link)
        print(f"Received audio file: {audio_path}")

        if audio_path and os.path.exists(audio_path):
            print("✅ Starting transcription and translation...")
            summary, translated_transcript = transcribe_and_translate(audio_path)

            if summary and translated_transcript:
                save_transcript_to_docx(summary, translated_transcript, "translated_transcript.docx")
            else:
                print("❌ Translation or transcription failed.")
        else:
            print("❌ Audio download failed or file not found.")
    else:
        print("❌ No YouTube link provided.")
